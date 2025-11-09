#!/usr/bin/env python3
"""
Comprehensive DOCX Format Extractor
提取 DOCX 文件的全面格式信息
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


def mm_to_pt(mm_val):
    """Convert millimeters to points"""
    return round(mm_val / 0.352778, 2)


def emu_to_mm(emu):
    """Convert EMU to millimeters"""
    if emu is None:
        return 0
    return round(emu / 36000, 2)


def twips_to_mm(twips):
    """Convert twips to millimeters"""
    if twips is None:
        return 0
    # 1 twip = 1/1440 inch, 1 inch = 25.4mm
    return round(twips / 1440 * 25.4, 2)


def twips_to_pt(twips):
    """Convert twips to points"""
    if twips is None:
        return 0
    return round(twips / 20, 1)


def extract_document_properties(doc):
    """Extract document properties"""
    print("\n【1. 文档属性】")
    print("="  * 60)

    try:
        core_props = doc.core_properties

        if core_props.title:
            print(f"标题: {core_props.title}")
        if core_props.author:
            print(f"作者: {core_props.author}")
        if core_props.subject:
            print(f"主题: {core_props.subject}")
        if core_props.comments:
            print(f"描述: {core_props.comments}")
        if core_props.created:
            print(f"创建时间: {core_props.created}")
        if core_props.modified:
            print(f"修改时间: {core_props.modified}")
        if core_props.category:
            print(f"类别: {core_props.category}")
        if core_props.keywords:
            print(f"关键字: {core_props.keywords}")

    except Exception as e:
        print(f"提取文档属性时出错: {e}")


def extract_page_setup(doc):
    """Extract page setup information"""
    print("\n【2. 页面设置】")
    print("=" * 60)

    try:
        section = doc.sections[0]

        # Page size
        page_width_mm = emu_to_mm(section.page_width)
        page_height_mm = emu_to_mm(section.page_height)
        print(f"页面尺寸: {page_width_mm}mm × {page_height_mm}mm")

        # Page orientation
        orientation = "横向" if section.orientation == 1 else "纵向"
        print(f"页面方向: {orientation}")

        # Margins
        print("\n页边距:")
        print(f"  上: {emu_to_mm(section.top_margin)}mm")
        print(f"  下: {emu_to_mm(section.bottom_margin)}mm")
        print(f"  左: {emu_to_mm(section.left_margin)}mm")
        print(f"  右: {emu_to_mm(section.right_margin)}mm")
        print(f"  页眉: {emu_to_mm(section.header_distance)}mm")
        print(f"  页脚: {emu_to_mm(section.footer_distance)}mm")

        # Gutters
        if hasattr(section, 'gutter'):
            print(f"  装订线: {emu_to_mm(section.gutter)}mm")

        print(f"\n总节数: {len(doc.sections)}")

    except Exception as e:
        print(f"提取页面设置时出错: {e}")


def extract_styles(doc):
    """Extract styles"""
    print("\n【3. 样式定义】")
    print("=" * 60)

    try:
        styles = doc.styles

        # Count styles by type
        para_styles = [s for s in styles if s.type == 1]  # PARAGRAPH
        char_styles = [s for s in styles if s.type == 2]  # CHARACTER
        table_styles = [s for s in styles if s.type == 3]  # TABLE

        print(f"段落样式数: {len(para_styles)}")
        print(f"字符样式数: {len(char_styles)}")
        print(f"表格样式数: {len(table_styles)}")
        print(f"总样式数: {len(list(styles))}")

        print("\n前10个段落样式:")
        for i, style in enumerate(para_styles[:10]):
            print(f"\n样式 {i+1}: {style.name}")
            if style.base_style:
                print(f"  基于: {style.base_style.name}")

            # Font
            if hasattr(style, 'font'):
                font = style.font
                if font.name:
                    print(f"  字体: {font.name}")
                if font.size:
                    print(f"  字号: {font.size.pt}磅")
                if font.bold:
                    print(f"  加粗: {font.bold}")
                if font.italic:
                    print(f"  斜体: {font.italic}")

    except Exception as e:
        print(f"提取样式时出错: {e}")


def get_paragraph_alignment(para):
    """Get paragraph alignment as string"""
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }
    return alignment_map.get(para.alignment, "未设置")


def extract_paragraph_format(para, para_num):
    """Extract format info for a paragraph"""
    info = []

    # Get text
    text = para.text.strip()
    if text:
        display_text = text[:50] + "..." if len(text) > 50 else text
        info.append(f"  内容: {display_text}")

    # Style
    if para.style:
        info.append(f"  样式: {para.style.name}")

    # Alignment
    if para.alignment is not None:
        info.append(f"  对齐: {get_paragraph_alignment(para)}")

    # Indentation
    pf = para.paragraph_format
    if pf.left_indent:
        info.append(f"  左缩进: {emu_to_mm(pf.left_indent)}mm")
    if pf.right_indent:
        info.append(f"  右缩进: {emu_to_mm(pf.right_indent)}mm")
    if pf.first_line_indent:
        info.append(f"  首行缩进: {emu_to_mm(pf.first_line_indent)}mm")

    # Spacing
    if pf.space_before:
        info.append(f"  段前间距: {twips_to_pt(pf.space_before)}磅")
    if pf.space_after:
        info.append(f"  段后间距: {twips_to_pt(pf.space_after)}磅")
    if pf.line_spacing:
        info.append(f"  行距: {pf.line_spacing}")

    # Run (character) formatting
    if para.runs:
        run = para.runs[0]
        run_info = []

        if run.font.name:
            run_info.append(f"字体={run.font.name}")
        if run.font.size:
            run_info.append(f"字号={run.font.size.pt}磅")
        if run.font.bold:
            run_info.append("加粗")
        if run.font.italic:
            run_info.append("斜体")
        if run.font.underline:
            run_info.append("下划线")
        if run.font.color and run.font.color.rgb:
            run_info.append(f"颜色=#{run.font.color.rgb}")

        if run_info:
            info.append(f"  字符格式: {', '.join(run_info)}")

    return info


def extract_content_formatting(doc):
    """Extract content formatting"""
    print("\n【4. 内容格式】")
    print("=" * 60)

    try:
        total_paras = len(doc.paragraphs)
        display_count = min(20, total_paras)

        for i, para in enumerate(doc.paragraphs[:display_count]):
            print(f"\n段落 {i+1}:")
            info = extract_paragraph_format(para, i+1)
            for line in info:
                print(line)

        if total_paras > display_count:
            print(f"\n... (还有 {total_paras - display_count} 个段落省略)")

        print(f"\n总段落数: {total_paras}")

    except Exception as e:
        print(f"提取内容格式时出错: {e}")


def extract_tables(doc):
    """Extract table information"""
    print("\n【5. 表格】")
    print("=" * 60)

    try:
        table_count = len(doc.tables)

        for i, table in enumerate(doc.tables):
            print(f"\n表格 {i+1}:")

            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0

            print(f"  行数: {rows}")
            print(f"  列数: {cols}")

            # Get first few cells content
            if rows > 0 and table.rows[0].cells:
                first_row_text = " | ".join([cell.text[:20] for cell in table.rows[0].cells[:5]])
                if first_row_text.strip():
                    print(f"  首行内容: {first_row_text}")

        print(f"\n总表格数: {table_count}")

    except Exception as e:
        print(f"提取表格时出错: {e}")


def extract_numbering(doc):
    """Extract numbering and lists"""
    print("\n【6. 编号和列表】")
    print("=" * 60)

    try:
        # Count paragraphs with numbering
        numbered_paras = []
        bulleted_paras = []

        for para in doc.paragraphs:
            if para.style.name.startswith('List'):
                if 'Number' in para.style.name:
                    numbered_paras.append(para)
                elif 'Bullet' in para.style.name:
                    bulleted_paras.append(para)

        print(f"编号列表项数: {len(numbered_paras)}")
        print(f"项目符号列表项数: {len(bulleted_paras)}")

        # Try to extract numbering info from XML
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part:
                print("\n文档包含编号定义")
        except:
            print("\n文档中没有编号定义")

    except Exception as e:
        print(f"提取编号时出错: {e}")


def extract_images(doc):
    """Extract image information"""
    print("\n【7. 图片和图形】")
    print("=" * 60)

    try:
        image_count = 0
        shape_count = 0

        # Count inline shapes (images)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        print(f"图片数量: {image_count}")

    except Exception as e:
        print(f"提取图片信息时出错: {e}")


def extract_headers_footers(doc):
    """Extract header and footer information"""
    print("\n【8. 页眉和页脚】")
    print("=" * 60)

    try:
        for i, section in enumerate(doc.sections):
            print(f"\n节 {i+1}:")

            # Header
            if section.header:
                header_text = " ".join([para.text for para in section.header.paragraphs]).strip()
                if header_text:
                    display_text = header_text[:50] + "..." if len(header_text) > 50 else header_text
                    print(f"  页眉: {display_text}")
                else:
                    print(f"  页眉: (空)")

            # Footer
            if section.footer:
                footer_text = " ".join([para.text for para in section.footer.paragraphs]).strip()
                if footer_text:
                    display_text = footer_text[:50] + "..." if len(footer_text) > 50 else footer_text
                    print(f"  页脚: {display_text}")
                else:
                    print(f"  页脚: (空)")

    except Exception as e:
        print(f"提取页眉页脚时出错: {e}")


def main():
    docx_path = "test.docx"

    print("=" * 60)
    print("    DOCX 格式全面提取分析")
    print(f"    文件: {docx_path}")
    print("    工具: python-docx")
    print("=" * 60)

    try:
        doc = Document(docx_path)

        # Extract all format information
        extract_document_properties(doc)
        extract_page_setup(doc)
        extract_styles(doc)
        extract_content_formatting(doc)
        extract_tables(doc)
        extract_numbering(doc)
        extract_images(doc)
        extract_headers_footers(doc)

        print("\n" + "=" * 60)
        print("    提取完成!")
        print("=" * 60)

    except FileNotFoundError:
        print(f"\n错误: 文件 '{docx_path}' 不存在")
        sys.exit(1)
    except Exception as e:
        print(f"\n错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
