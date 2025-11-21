"""
文档生成器 - 将解析后的内容和样式结合生成 docx 文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.table import _Cell
from .styles import StyleManager
import os
import re


CITATION_PATTERN = re.compile(r'\[(\d+)\]')
WORD_JOINER = '\u2060'
SPECIAL_SECTION_BOOKMARKS = {
    'references': '_Section_References',
    'acknowledgements': '_Section_Acknowledgements',
    'appendix': '_Section_Appendix'
}


class ThesisGenerator:
    """论文文档生成器"""

    def __init__(self, style_manager):
        """
        初始化生成器
        :param style_manager: 样式管理器实例
        """
        self.style_manager = style_manager
        self._reset_document()

    def _reset_document(self):
        """重新创建文档，确保每次生成都是干净的"""
        self.doc = Document()
        self.bookmark_id = 0  # 书签ID计数器
        self.reference_targets = {}
        self.references_data = []
        self.reference_backlinks = {}
        self._setup_document()

    def _setup_document(self):
        """设置文档基本属性"""
        for section in self.doc.sections:
            self._apply_section_layout(section)
        self._apply_normal_style_defaults()

    def _apply_section_layout(self, section):
        """根据配置设置节的页面属性"""
        document_settings = self.style_manager.get_document_settings()
        margins = document_settings.get('margins', {})

        # 设置纸张大小
        page_size = document_settings.get('page_size', 'A4')
        if page_size == 'A4':
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        elif page_size == 'Letter':
            from docx.shared import Inches
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # 设置页边距
        section.top_margin = Cm(margins.get('top', 2.0))
        section.bottom_margin = Cm(margins.get('bottom', 2.0))
        section.left_margin = Cm(margins.get('left', 2.5))
        section.right_margin = Cm(margins.get('right', 2.0))
        section.header_distance = Cm(document_settings.get('header_distance', 1.27))
        section.footer_distance = Cm(document_settings.get('footer_distance', 1.27))
        section.gutter = Cm(document_settings.get('gutter', 0))

    def _apply_normal_style_defaults(self):
        """统一设置 Normal 样式，避免 Word 默认字体干扰"""
        fonts = self.style_manager.get_fonts()
        paragraph_style = self.style_manager.get_paragraph_style()
        english_font = fonts.get('english', 'Times New Roman')
        chinese_font = fonts.get('chinese', '宋体')

        normal_style = self.doc.styles['Normal']
        normal_style.font.name = english_font
        normal_style.font.size = Pt(paragraph_style.get('size', 12))

        rPr = normal_style.element.get_or_add_rPr()
        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), chinese_font)
        rFonts.set(qn('w:ascii'), english_font)
        rFonts.set(qn('w:hAnsi'), english_font)

        normal_format = normal_style.paragraph_format
        normal_format.space_before = Pt(0)
        normal_format.space_after = Pt(0)
        normal_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        self._apply_math_defaults()

    def _apply_math_defaults(self):
        """设置全局数学公式默认字体等属性，避免Word回退到Cambria Math"""
        try:
            formula_style = self.style_manager.get_formula_style()
        except KeyError:
            return

        font_name = formula_style.get('font', 'Times New Roman')
        settings = self.doc.settings
        settings_element = settings.element

        math_pr = settings_element.find(qn('m:mathPr'))
        if math_pr is None:
            math_pr = OxmlElement('m:mathPr')
            settings_element.append(math_pr)

        math_font = math_pr.find(qn('m:mathFont'))
        if math_font is None:
            math_font = OxmlElement('m:mathFont')
            math_pr.append(math_font)

        math_font.set(qn('m:val'), font_name)
        w_val_attr = qn('w:val')
        if w_val_attr in math_font.attrib:
            del math_font.attrib[w_val_attr]

    def _add_configured_section(self):
        """新增节并沿用统一的页面设置"""
        section = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        self._apply_section_layout(section)
        return section
    
    def _create_special_section(self, header_title, page_number_config):
        """
        创建用于附加章节（参考文献/致谢/附录）的节，设置页码与页眉
        """
        section = self._add_configured_section()
        if page_number_config:
            self._apply_page_number_settings(section, page_number_config)
        else:
            self._ensure_continuous_page_numbering(section)
        self._set_header(header_title, section)
        return section

    def _clear_block_paragraphs(self, block):
        """清空页眉/页脚已有内容"""
        for para in list(block.paragraphs):
            p_element = para._element
            p_element.getparent().remove(p_element)

    def _prepare_reference_targets(self, references):
        """为参考文献生成对应的书签映射"""
        self.references_data = references or []
        self.reference_targets = {}
        for idx in range(len(self.references_data)):
            bookmark_name = f'_Reference_{idx + 1}'
            self.reference_targets[idx + 1] = {'bookmark': bookmark_name}

    def _apply_page_number_settings(self, section, config):
        """根据配置为节设置页码格式"""
        if not config:
            return

        self._configure_page_number_type(section, config)
        footer = section.footer
        footer.is_linked_to_previous = False
        self._clear_block_paragraphs(footer)
        footer_para = footer.add_paragraph()

        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        footer_para.alignment = alignment_map.get(config.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)

        page_run = footer_para.add_run()
        font_name = config.get('font', 'Times New Roman')
        page_run.font.name = font_name
        page_run.font.size = Pt(config.get('size', 9))
        page_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        field_instr = self._build_page_field_instruction(config.get('style', 'arabic'))
        self._add_field_code(page_run, field_instr)

    def _configure_page_number_type(self, section, config):
        """设置节的起始页码与格式"""
        sectPr = section._sectPr
        pg_num = sectPr.find(qn('w:pgNumType'))
        if pg_num is None:
            pg_num = OxmlElement('w:pgNumType')
            sectPr.append(pg_num)

        if 'start_from' in config:
            pg_num.set(qn('w:start'), str(config['start_from']))

        fmt_value = self._map_page_number_format(config.get('style', 'arabic'))
        if fmt_value:
            pg_num.set(qn('w:fmt'), fmt_value)

    def _map_page_number_format(self, style_name):
        """将配置中的样式映射为 Word 的格式值"""
        mapping = {
            'roman': 'lowerRoman',
            'ROMAN': 'upperRoman',
            'arabic': 'decimal'
        }
        return mapping.get(style_name, 'decimal')

    def _build_page_field_instruction(self, style_name):
        """生成 PAGE 字段指令"""
        if style_name == 'roman':
            return 'PAGE \\* roman'
        if style_name == 'ROMAN':
            return 'PAGE \\* ROMAN'
        return 'PAGE'

    def _add_field_code(self, run, instruction):
        """在 run 中插入字段指令"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instruction

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _get_next_bookmark_id(self):
        """获取下一个书签ID"""
        self.bookmark_id += 1
        return self.bookmark_id

    def _add_bookmark_to_paragraph(self, paragraph, bookmark_name):
        """
        为段落添加书签
        :param paragraph: 段落对象
        :param bookmark_name: 书签名称
        """
        bookmark_id = self._get_next_bookmark_id()

        # 创建书签开始标记
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), bookmark_name)

        # 创建书签结束标记
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))

        # 插入到段落
        p_element = paragraph._element
        if len(p_element) > 0:
            p_element.insert(0, bookmark_start)
            p_element.append(bookmark_end)
        else:
            p_element.append(bookmark_start)
            p_element.append(bookmark_end)

    def _get_special_section_bookmark(self, section_key):
        """获取特殊章节对应的书签名称"""
        return SPECIAL_SECTION_BOOKMARKS.get(section_key)

    def _resolve_title_text(self, title_cfg, default_text):
        """根据配置解析标题文本"""
        if not title_cfg:
            return default_text
        return title_cfg.get('text', default_text)

    def _get_special_section_title(self, section_key):
        """根据样式配置获取特殊章节标题"""
        if section_key == 'references':
            config = self.style_manager.get_references_style() or {}
            title_cfg = config.get('title', {})
            default = '参考文献'
        elif section_key == 'acknowledgements':
            config = self.style_manager.get_acknowledgement_style() or {}
            title_cfg = config.get('title', {})
            default = '致  谢'
        elif section_key == 'appendix':
            config = self.style_manager.get_appendix_style() or {}
            title_cfg = config.get('title', {})
            default = '附  录'
        else:
            return ''
        return self._resolve_title_text(title_cfg, default)

    def _format_section_title(self, section_number, title_text):
        """兼容旧逻辑，当前直接返回标题文本"""
        return title_text or ''

    def _get_last_chapter_number(self, chapters):
        """获取正文中最后一个章节号（若缺失则返回章节总数）"""
        last_number = 0
        for chapter in chapters:
            num = chapter.get('number')
            if num is None:
                continue
            try:
                last_number = int(str(num))
            except (ValueError, TypeError):
                continue
        if last_number:
            return last_number
        return len(chapters)

    def _create_standard_hyperlink(self, paragraph, text, bookmark_name, font_size=10.5):
        """
        创建符合Word标准的内部超链接
        :param paragraph: 段落对象
        :param text: 超链接文本
        :param bookmark_name: 目标书签名称
        :param font_size: 字体大小（磅）
        """
        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        hyperlink.set(qn('w:history'), '1')

        # 创建运行元素
        run_element = OxmlElement('w:r')

        # 创建运行属性
        run_props = OxmlElement('w:rPr')

        # 关键：设置超链接样式
        style_element = OxmlElement('w:rStyle')
        style_element.set(qn('w:val'), 'Hyperlink')
        run_props.append(style_element)

        # 设置字体（中英文分离处理）
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Times New Roman')
        fonts.set(qn('w:eastAsia'), '宋体')
        fonts.set(qn('w:hAnsi'), 'Times New Roman')
        run_props.append(fonts)

        # 设置字体大小（磅转为半磅）
        size_element = OxmlElement('w:sz')
        size_element.set(qn('w:val'), str(int(font_size * 2)))
        run_props.append(size_element)

        # 设置颜色（黑色文本）
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        run_props.append(color)

        run_element.append(run_props)

        # 添加文本内容
        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)

        hyperlink.append(run_element)
        paragraph._element.append(hyperlink)

    def _add_pageref_field(self, run, bookmark_name):
        """
        添加PAGEREF字段来自动获取页码
        :param run: 运行对象
        :param bookmark_name: 书签名称
        """
        # 创建字段开始标记
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        # 创建指令文本
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'PAGEREF {bookmark_name} \\h'  # \h 表示超链接格式

        # 创建字段结束标记
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # 将字段元素添加到运行中
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _add_tab_stop(self, paragraph, position_cm=16.0, alignment='right', leader=None):
        """
        为段落添加制表位
        :param paragraph: 段落对象
        :param position_cm: 制表位位置（厘米）
        :param alignment: 制表位对齐方式（left/center/right）
        :param leader: 前导符类型（如 'dot'）
        """
        pPr = paragraph._element.get_or_add_pPr()
        tabs = pPr.find(qn('w:tabs'))
        if tabs is None:
            tabs = OxmlElement('w:tabs')
            pPr.append(tabs)

        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), alignment)
        if leader:
            tab.set(qn('w:leader'), leader)
        tab.set(qn('w:pos'), str(int(position_cm * 567)))
        tabs.append(tab)

    def generate(self, content, output_path, include_toc=True):
        """
        生成完整的论文文档
        :param content: 解析后的内容结构
        :param output_path: 输出文件路径
        :param include_toc: 是否包含目录
        """
        import os
        import glob

        # 每次生成都重新初始化文档
        self._reset_document()

        # 删除输出目录中的所有旧docx文件
        output_dir = os.path.dirname(output_path)
        if output_dir and os.path.exists(output_dir):
            old_docx_files = glob.glob(os.path.join(output_dir, '*.docx'))
            for old_file in old_docx_files:
                try:
                    os.remove(old_file)
                    print(f"已删除旧文件: {old_file}")
                except Exception as e:
                    print(f"删除文件失败 {old_file}: {str(e)}")

        chapters = content.get('chapters', [])
        has_chapters = bool(chapters)
        include_toc = include_toc and has_chapters
        has_cn_abstract = bool(content.get('abstract'))
        has_en_abstract = bool(content.get('abstract_en') and content['abstract_en'].get('content'))
        has_abstract = has_cn_abstract or has_en_abstract
        references = content.get('references', [])
        has_references = bool(references)
        acknowledgements = content.get('acknowledgements', [])
        has_ack = bool(acknowledgements)
        appendix_entries = content.get('appendix', [])
        has_appendix = bool(appendix_entries)
        if has_references:
            self._prepare_reference_targets(references)
        else:
            self.reference_targets = {}
            self.references_data = []

        section_number_map = {}
        current_section_number = self._get_last_chapter_number(chapters)
        for key, enabled in (
            ('references', has_references),
            ('acknowledgements', has_ack),
            ('appendix', has_appendix)
        ):
            if not enabled:
                continue
            current_section_number += 1
            section_number_map[key] = current_section_number

        special_sections_for_toc = []
        if include_toc:
            for key in ('references', 'acknowledgements', 'appendix'):
                if key not in section_number_map:
                    continue
                special_sections_for_toc.append({
                    'title': self._format_section_title(
                        section_number_map[key],
                        self._get_special_section_title(key)
                    ),
                    'bookmark': self._get_special_section_bookmark(key)
                })

        base_section = self.doc.sections[0]
        body_section = None
        body_page_config = self.style_manager.get_page_number_config('body') or {}

        # 生成中文摘要部分
        if has_cn_abstract:
            self._apply_page_number_settings(
                base_section,
                self.style_manager.get_page_number_config('abstract')
            )
            self._generate_abstract(content['abstract'])
            # 摘要后添加分节符
            self.doc.add_page_break()

        # 生成英文摘要部分
        if has_en_abstract:
            if not has_cn_abstract:
                self._apply_page_number_settings(
                    base_section,
                    self.style_manager.get_page_number_config('abstract')
                )
            self._generate_abstract_en(content['abstract_en'])

        if include_toc:
            if has_abstract:
                toc_section = self._add_configured_section()
            else:
                toc_section = base_section
                if not (has_cn_abstract or has_en_abstract):
                    self._apply_page_number_settings(
                        toc_section,
                        self.style_manager.get_page_number_config('toc')
                    )
            if toc_section is not base_section:
                self._apply_page_number_settings(
                    toc_section,
                    self.style_manager.get_page_number_config('toc')
                )
            self._generate_toc(chapters, special_sections_for_toc)
            body_section = self._add_configured_section()
            self._apply_page_number_settings(
                body_section,
                body_page_config
            )
        else:
            if has_abstract:
                body_section = self._add_configured_section()
                self._apply_page_number_settings(
                    body_section,
                    body_page_config
                )
            else:
                body_section = base_section
                self._apply_page_number_settings(
                    body_section,
                    body_page_config
                )

        if has_chapters:
            self._generate_body(content['title'], chapters, body_section)

        if has_references:
            references_header = self._format_section_title(
                section_number_map.get('references'),
                self._get_special_section_title('references')
            ) or '参考文献'
            self._create_special_section(references_header, None)
            self._generate_references(references, section_number_map.get('references'))

        if has_ack:
            ack_header = self._format_section_title(
                section_number_map.get('acknowledgements'),
                self._get_special_section_title('acknowledgements')
            ) or '致谢'
            self._create_special_section(ack_header, None)
            self._generate_acknowledgements(acknowledgements, section_number_map.get('acknowledgements'))

        if has_appendix:
            appendix_header = self._format_section_title(
                section_number_map.get('appendix'),
                self._get_special_section_title('appendix')
            ) or '附录'
            self._create_special_section(appendix_header, None)
            self._generate_appendix(appendix_entries, section_number_map.get('appendix'))

        # 保存文档
        self.doc.save(output_path)
        print(f"新文档已生成: {output_path}")

    def _generate_abstract(self, abstract_data):
        """
        生成摘要部分
        :param abstract_data: 摘要数据
        """
        fonts = self.style_manager.get_fonts()
        english_font = fonts.get('english', 'Times New Roman')

        # 标题
        title_style = self.style_manager.get_abstract_title_style()
        title_text = title_style.get('text', '摘 要')
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title_text)
        self.style_manager.apply_run_style(title_run, title_style)
        self._apply_title_paragraph_format(title_para, title_style)

        # 内容
        content_style = self.style_manager.get_abstract_content_style()
        for para_text in abstract_data.get('content', []):
            if not para_text:
                continue
            para = self.doc.add_paragraph()
            run = para.add_run()
            self.style_manager.set_mixed_font(
                run,
                para_text,
                chinese_font=content_style.get('font', '宋体'),
                english_font=content_style.get('english_font', english_font),
                size=content_style.get('size', 12),
                bold=content_style.get('bold', False)
            )
            self.style_manager.apply_paragraph_style(para, content_style)

        # 关键词
        keywords = abstract_data.get('keywords', [])
        if keywords:
            kw_style = self.style_manager.get_abstract_keywords_style()
            kw_para = self.doc.add_paragraph()

            # 段落样式（需要尺寸配合悬挂缩进）
            kw_style_para = dict(kw_style)
            kw_style_para.setdefault('size', kw_style.get('content_size', 12))
            self.style_manager.apply_paragraph_style(kw_para, kw_style_para)

            # 标签
            label_text = kw_style.get('label_text', '关键词：')
            label_run = kw_para.add_run(label_text)
            label_run.font.name = kw_style.get('label_font', '宋体')
            label_run._element.rPr.rFonts.set(qn('w:eastAsia'), kw_style.get('label_font', '宋体'))
            label_run.font.size = Pt(kw_style.get('label_size', kw_style_para['size']))
            label_run.font.bold = kw_style.get('label_bold', True)

            # 内容
            keywords_text = kw_style.get('separator', '；').join(keywords)
            content_run = kw_para.add_run()
            self.style_manager.set_mixed_font(
                content_run,
                keywords_text,
                chinese_font=kw_style.get('content_font', '宋体'),
                english_font=kw_style.get('content_font', english_font),
                size=kw_style.get('content_size', 12)
            )

    def _generate_abstract_en(self, abstract_data):
        """
        生成英文摘要部分
        :param abstract_data: 英文摘要数据
        """
        en_cfg = self.style_manager.config.get('abstract_en', {})

        # 标题
        title_cfg = self.style_manager.get_abstract_en_title_style()
        title_text = title_cfg.get('text', 'ABSTRACT')
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title_text)
        self.style_manager.apply_run_style(title_run, title_cfg)
        self._apply_title_paragraph_format(title_para, title_cfg)

        # 内容
        content_cfg = self.style_manager.get_abstract_en_content_style()
        for para_text in abstract_data.get('content', []):
            if not para_text:
                continue
            para = self.doc.add_paragraph()
            run = para.add_run(para_text)
            self.style_manager.apply_run_style(run, content_cfg)
            self.style_manager.apply_paragraph_style(para, content_cfg)

        # 关键词
        keywords = abstract_data.get('keywords', [])
        if keywords:
            kw_cfg = self.style_manager.get_abstract_en_keywords_style()
            kw_para = self.doc.add_paragraph()
            kw_style_para = dict(kw_cfg)
            kw_style_para.setdefault('size', kw_cfg.get('content_size', 12))
            self.style_manager.apply_paragraph_style(kw_para, kw_style_para)

            label_text = kw_cfg.get('label_text', 'Key Words:')
            label_run = kw_para.add_run(label_text)
            label_run.font.name = kw_cfg.get('label_font', 'Times New Roman')
            label_run.font.size = Pt(kw_cfg.get('label_size', kw_style_para['size']))
            label_run.font.bold = kw_cfg.get('label_bold', True)

            keywords_text = kw_cfg.get('separator', ';').join(
                [kw.capitalize() if kw_cfg.get('capitalize_each') else kw for kw in keywords]
            )
            content_run = kw_para.add_run(' ' + keywords_text if keywords_text else '')
            content_run.font.name = kw_cfg.get('content_font', 'Times New Roman')
            content_run.font.size = Pt(kw_cfg.get('content_size', 12))

    def _generate_toc(self, chapters, special_sections=None):
        """
        生成可点击跳转的目录（带自动页码）
        :param chapters: 章节列表
        :param special_sections: 追加的特殊章节条目（参考文献/致谢/附录）
        """
        special_sections = special_sections or []
        toc_cfg = self.style_manager.get_toc_config()
        title_cfg = toc_cfg.get('title', {})
        entry_cfg = toc_cfg.get('entry', {})
        page_cfg = toc_cfg.get('page_number', {})
        body_cfg = self.style_manager.config.get('body', {})
        h1_fmt = body_cfg.get('heading1', {}).get('number_format', '第{chapter}章')

        # 标题
        toc_title = self.doc.add_paragraph()
        toc_text = title_cfg.get('text', '目 录')
        toc_run = toc_title.add_run(toc_text)
        self.style_manager.apply_run_style(toc_run, title_cfg or {})
        self._apply_title_paragraph_format(toc_title, title_cfg or {})

        entry_font = entry_cfg.get('font', '宋体')
        entry_size = entry_cfg.get('size', 12)
        indent_rules = entry_cfg.get('indent_rules', {})
        level1_indent = Pt(indent_rules.get('level1_chars', 0) * entry_size)
        level2_indent = Pt(indent_rules.get('level2_chars', 0) * entry_size)
        level3_indent = Pt(indent_rules.get('level3_chars', 0) * entry_size)
        page_font = page_cfg.get('font', 'Times New Roman')
        page_size = page_cfg.get('size', 10.5)

        # 目录条目
        for chapter in chapters:
            chapter_num = chapter.get('number', 1)
            chapter_title = chapter.get('title', '')
            chapter_cn = self._arabic_to_chinese(chapter_num)
            bookmark_name = f'_Chapter_{chapter_num}'

            toc_p = self.doc.add_paragraph()
            toc_p.paragraph_format.left_indent = level1_indent
            self._add_tab_stop(toc_p, position_cm=16.0, alignment='right', leader='dot')

            entry_text = h1_fmt.format(chapter=chapter_num, chapter_cn=chapter_cn) + f' {chapter_title}'
            self._create_standard_hyperlink(toc_p, entry_text, bookmark_name)
            toc_p.add_run('\t')
            page_run = toc_p.add_run()
            page_run.font.name = page_font
            page_run.font.size = Pt(page_size)
            self._add_pageref_field(page_run, bookmark_name)

            for item in chapter.get('content', []):
                if item['type'] == 'heading2':
                    h2_bookmark = self._make_heading_bookmark(item["number"])
                    h2_p = self.doc.add_paragraph()
                    h2_p.paragraph_format.left_indent = level2_indent
                    self._add_tab_stop(h2_p, position_cm=16.0, alignment='right', leader='dot')
                    h2_text = f'{item["number"]} {item["text"]}'
                    self._create_standard_hyperlink(h2_p, h2_text, h2_bookmark)
                    h2_p.add_run('\t')
                    h2_page_run = h2_p.add_run()
                    h2_page_run.font.name = page_font
                    h2_page_run.font.size = Pt(page_size)
                    self._add_pageref_field(h2_page_run, h2_bookmark)

                elif item['type'] == 'heading3':
                    h3_bookmark = self._make_heading_bookmark(item["number"])
                    h3_p = self.doc.add_paragraph()
                    h3_p.paragraph_format.left_indent = level3_indent
                    self._add_tab_stop(h3_p, position_cm=16.0, alignment='right', leader='dot')
                    h3_text = f'{item["number"]} {item["text"]}'
                    self._create_standard_hyperlink(h3_p, h3_text, h3_bookmark)
                    h3_p.add_run('\t')
                    h3_page_run = h3_p.add_run()
                    h3_page_run.font.name = page_font
                    h3_page_run.font.size = Pt(page_size)
                    self._add_pageref_field(h3_page_run, h3_bookmark)

        for section in special_sections:
            title = section.get('title')
            bookmark_name = section.get('bookmark')
            if not title or not bookmark_name:
                continue
            section_para = self.doc.add_paragraph()
            section_para.paragraph_format.left_indent = level1_indent
            self._add_tab_stop(section_para, position_cm=16.0, alignment='right', leader='dot')
            self._create_standard_hyperlink(section_para, title, bookmark_name)
            section_para.add_run('\t')
            page_run = section_para.add_run()
            page_run.font.name = page_font
            page_run.font.size = Pt(page_size)
            self._add_pageref_field(page_run, bookmark_name)

    def _generate_body(self, title, chapters, section):
        """
        生成正文部分
        :param title: 论文标题（用于页眉）
        :param chapters: 章节列表
        :param section: 正文所在的节
        """
        # 设置正文页眉
        self._set_header(title, section)

        # 生成各章节
        for chapter_idx, chapter in enumerate(chapters):
            self._generate_chapter(chapter, chapter_idx)

    def _generate_references(self, references, section_number=None):
        """
        生成参考文献部分
        :param references: 参考文献列表
        """
        if not references:
            return

        ref_config = self.style_manager.get_references_style()
        title_cfg = ref_config.get('title', {})
        entry_cfg = ref_config.get('entry', {})
        number_cfg = ref_config.get('number', {})
        fonts = self.style_manager.get_fonts()
        english_font = fonts.get('english', 'Times New Roman')

        # 标题
        title_text = self._resolve_title_text(title_cfg, '参考文献')
        display_title = self._format_section_title(section_number, title_text)
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(display_title)
        if title_cfg:
            self.style_manager.apply_run_style(title_run, title_cfg)
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if 'alignment' in title_cfg:
            title_para.alignment = alignment_map.get(title_cfg['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        if 'space_before' in title_cfg:
            title_para.paragraph_format.space_before = Pt(title_cfg.get('space_before', 0))
        if 'space_after' in title_cfg:
            title_para.paragraph_format.space_after = Pt(title_cfg.get('space_after', 0))

        bookmark_name = self._get_special_section_bookmark('references')
        if bookmark_name:
            self._add_bookmark_to_paragraph(title_para, bookmark_name)

        # 条目
        entry_font = entry_cfg.get('font', '宋体')
        entry_size = entry_cfg.get('size', 10.5)
        number_font = number_cfg.get('font', entry_font)
        number_bold = number_cfg.get('bold', False)

        for idx, ref in enumerate(references, 1):
            para = self.doc.add_paragraph()
            self.style_manager.apply_paragraph_style(para, entry_cfg)

            if 'prefix' in number_cfg or 'suffix' in number_cfg:
                number_text = f'{number_cfg.get("prefix", "")}{idx}{number_cfg.get("suffix", "")}'
            else:
                left_bracket, right_bracket = number_cfg.get('brackets', ['[', ']'])
                number_text = f'{left_bracket}{idx}{right_bracket}'

            number_run = para.add_run(number_text)
            number_run.font.name = number_font
            number_run.font.size = Pt(entry_size)
            number_run.font.bold = number_bold
            number_run._element.rPr.rFonts.set(qn('w:eastAsia'), entry_font)

            text = self._sanitize_reference_text(ref.get('text', ''))
            ref['text'] = text
            detail_text = f' {text}' if text else ''
            backlink_name = self.reference_backlinks.get(idx)
            if detail_text:
                if backlink_name:
                    self._add_internal_reference_link(
                        para,
                        detail_text,
                        backlink_name,
                        entry_font,
                        english_font,
                        entry_size
                    )
                else:
                    detail_run = para.add_run()
                    self.style_manager.set_mixed_font(
                        detail_run,
                        detail_text,
                        chinese_font=entry_font,
                        english_font=english_font,
                        size=entry_size
                    )

            target = self.reference_targets.get(idx)
            if target:
                self._add_bookmark_to_paragraph(para, target['bookmark'])

    def _sanitize_reference_text(self, text):
        """移除 URL、统一标点并控制换行"""
        if not text:
            return ''

        sanitized = text.strip()
        sanitized = self._remove_reference_urls(sanitized)
        sanitized = self._normalize_reference_punctuation(sanitized)
        sanitized = self._remove_space_before_punctuation(sanitized)
        sanitized = self._collapse_reference_whitespace(sanitized)
        sanitized = self._ensure_reference_spacing(sanitized)
        sanitized = self._collapse_reference_whitespace(sanitized)
        sanitized = self._protect_reference_sequences(sanitized)
        sanitized = sanitized.strip()
        sanitized = re.sub(r'\s+\.', '.', sanitized)
        if sanitized and not sanitized.endswith('.'):
            sanitized = f'{sanitized}.'
        return sanitized

    def _remove_reference_urls(self, text):
        """去除参考文献中的 http/https 链接"""
        return re.sub(r'https?://\S+', '', text, flags=re.IGNORECASE)

    def _normalize_reference_punctuation(self, text):
        """将标点统一为英文半角字符"""
        replacements = {
            '……': '...'
        }
        for original, repl in replacements.items():
            text = text.replace(original, repl)

        punctuation_map = {
            '，': ',',
            '。': '.',
            '．': '.',
            '、': ',',
            '；': ';',
            '：': ':',
            '？': '?',
            '！': '!',
            '（': '(',
            '）': ')',
            '【': '[',
            '】': ']',
            '《': '<',
            '》': '>',
            '“': '"',
            '”': '"',
            '‘': "'",
            '’': "'",
            '—': '-',
            '－': '-',
            '～': '~',
            '·': '-',
            '｜': '|'
        }
        translation_table = str.maketrans(punctuation_map)
        return text.translate(translation_table)

    def _remove_space_before_punctuation(self, text):
        """清理标点前多余的空格"""
        return re.sub(r'\s+([,.;:?!\)\]\}])', r'\1', text)

    def _ensure_reference_spacing(self, text):
        """在标点后插入空格（- 和句末句号除外）"""
        if not text:
            return ''

        punctuation_pattern = re.compile(r'([,;:?!\)\]\}])(?=[^\s,.;:?!\)\]\}])')
        text = punctuation_pattern.sub(r'\1 ', text)

        period_pattern = re.compile(r'\.(?=[^\s,.;:?!\)\]\}])')
        text = period_pattern.sub('. ', text)
        return text

    def _collapse_reference_whitespace(self, text):
        """压缩多余空白"""
        return re.sub(r'\s+', ' ', text)

    def _protect_reference_sequences(self, text):
        """为日期、页码等连字符片段添加不可断开控制"""
        def wrap(match):
            segment = match.group(0)
            return segment.replace('-', f'{WORD_JOINER}-{WORD_JOINER}')

        return re.sub(r'\d+(?:-\d+)+', wrap, text)

    def _ensure_continuous_page_numbering(self, section):
        """移除新节的起始页码设置，保持与上一节连续"""
        sectPr = section._sectPr
        pg_num = sectPr.find(qn('w:pgNumType'))
        if pg_num is not None:
            start_attr = pg_num.get(qn('w:start'))
            if start_attr is not None:
                pg_num.attrib.pop(qn('w:start'))
        # 确保页脚延用上一节的页码字段
        section.footer.is_linked_to_previous = True

    def _make_heading_bookmark(self, number_text):
        """将标题编号转换为安全的书签名称"""
        safe = re.sub(r'[^0-9A-Za-z]+', '_', number_text)
        if not safe.strip('_'):
            safe = ''.join(str(ord(ch)) for ch in number_text)
        safe = safe.strip('_') or '0'
        return f'_Heading_{safe}'

    def _arabic_to_chinese(self, number):
        """阿拉伯数字转中文（1-99，满足当前需求）"""
        mapping = {0: '零', 1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
                   6: '六', 7: '七', 8: '八', 9: '九', 10: '十'}
        if number <= 10:
            return mapping[number]
        if number < 20:
            return '十' + mapping[number - 10] if number - 10 else '十'
        tens, ones = divmod(number, 10)
        tens_part = mapping.get(tens, str(tens)) + '十'
        ones_part = mapping.get(ones, '') if ones else ''
        return tens_part + ones_part

    def _generate_acknowledgements(self, paragraphs, section_number=None):
        """生成致谢部分"""
        section_style = self.style_manager.get_acknowledgement_style()
        self._render_custom_section(
            section_style,
            paragraphs,
            default_title='致  谢',
            section_number=section_number,
            bookmark_name=self._get_special_section_bookmark('acknowledgements')
        )

    def _generate_appendix(self, paragraphs, section_number=None):
        """生成附录部分"""
        section_style = self.style_manager.get_appendix_style()
        self._render_custom_section(
            section_style,
            paragraphs,
            default_title='附  录',
            section_number=section_number,
            bookmark_name=self._get_special_section_bookmark('appendix')
        )

    def _render_custom_section(self, section_style, paragraphs, default_title, section_number=None, bookmark_name=None):
        """渲染自定义章节（致谢/附录）"""
        if not paragraphs:
            return

        section_style = section_style or {}
        title_cfg = section_style.get('title', {})
        content_cfg = section_style.get('content', {})

        title_text = self._resolve_title_text(title_cfg, default_title)
        display_title = self._format_section_title(section_number, title_text)
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(display_title)
        if title_cfg:
            self.style_manager.apply_run_style(title_run, title_cfg)
        self._apply_title_paragraph_format(title_para, title_cfg)

        if bookmark_name:
            self._add_bookmark_to_paragraph(title_para, bookmark_name)

        fonts = self.style_manager.get_fonts()
        english_font = fonts.get('english', 'Times New Roman')
        content_font = content_cfg.get('font', '宋体')
        content_size = content_cfg.get('size', 12)

        for para_text in paragraphs:
            if not para_text:
                continue
            para = self.doc.add_paragraph()
            run = para.add_run()
            self.style_manager.set_mixed_font(
                run,
                para_text,
                chinese_font=content_font,
                english_font=english_font,
                size=content_size
            )
            self.style_manager.apply_paragraph_style(para, content_cfg or {})

    def _apply_title_paragraph_format(self, paragraph, title_cfg):
        """应用标题段落的对齐与间距"""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if title_cfg.get('alignment'):
            paragraph.alignment = alignment_map.get(title_cfg['alignment'], WD_ALIGN_PARAGRAPH.CENTER)
        if 'space_before' in title_cfg:
            paragraph.paragraph_format.space_before = Pt(title_cfg.get('space_before', 0))
        if 'space_after' in title_cfg:
            paragraph.paragraph_format.space_after = Pt(title_cfg.get('space_after', 0))

    def _generate_chapter(self, chapter, chapter_idx):
        """
        生成章节
        :param chapter: 章节数据
        :param chapter_idx: 章节索引
        """
        # 一级标题
        h1_style = self.style_manager.get_heading_style(1)
        h1_para = self.doc.add_paragraph()

        chapter_num = chapter.get('number', chapter_idx + 1)
        chapter_cn = self._arabic_to_chinese(chapter_num)
        number_format = h1_style.get('number_format', '第{chapter}章')
        heading_text = number_format.format(chapter=chapter_num, chapter_cn=chapter_cn)

        num_run = h1_para.add_run(f'{heading_text} ')
        num_run.font.name = h1_style.get('number_font', h1_style.get('font', '黑体'))
        num_run._element.rPr.rFonts.set(qn('w:eastAsia'), h1_style.get('font', '黑体'))
        num_run.font.size = Pt(h1_style['size'])
        num_run.font.bold = h1_style.get('bold', False)

        title_run = h1_para.add_run(chapter['title'])
        title_run.font.name = h1_style['font']
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), h1_style['font'])
        title_run.font.size = Pt(h1_style['size'])
        title_run.font.bold = h1_style.get('bold', False)

        self._apply_title_paragraph_format(h1_para, h1_style)

        # 🔑 关键：为一级标题添加书签
        bookmark_name = f'_Chapter_{chapter_num}'
        self._add_bookmark_to_paragraph(h1_para, bookmark_name)

        # 添加章节内容
        for item in chapter.get('content', []):
            if item['type'] == 'paragraph':
                self._add_paragraph(item['text'])
            elif item['type'] == 'heading2':
                self._add_heading2(item['number'], item['text'])
            elif item['type'] == 'heading3':
                self._add_heading3(item['number'], item['text'])
            elif item['type'] == 'figure':
                self._add_figure(item)
            elif item['type'] == 'table':
                self._add_table(item)
            elif item['type'] == 'formula':
                self._add_formula(item)

    def _add_paragraph(self, text):
        """
        添加正文段落
        :param text: 段落文本
        """
        para_style = self.style_manager.get_paragraph_style()
        para = self.doc.add_paragraph()

        fonts = self.style_manager.get_fonts()
        english_font = fonts.get('english', 'Times New Roman')

        self._add_text_with_citations(
            para,
            text,
            chinese_font=para_style['font'],
            english_font=english_font,
            font_size=para_style['size'],
            bold=para_style.get('bold', False)
        )

        # 应用段落样式
        self.style_manager.apply_paragraph_style(para, para_style)

    def _add_heading2(self, number, text):
        """
        添加二级标题
        :param number: 标题编号
        :param text: 标题文本
        """
        h2_style = self.style_manager.get_heading_style(2)
        para = self.doc.add_paragraph()

        # 编号（Times New Roman）
        num_run = para.add_run(f'{number} ')
        num_run.font.name = h2_style.get('number_font', 'Times New Roman')
        num_run.font.size = Pt(h2_style['size'])
        num_run.font.bold = h2_style.get('bold', False)

        # 标题文本（宋体）
        text_run = para.add_run(text)
        text_run.font.name = h2_style['font']
        text_run._element.rPr.rFonts.set(qn('w:eastAsia'), h2_style['font'])
        text_run.font.size = Pt(h2_style['size'])
        text_run.font.bold = h2_style.get('bold', False)

        # 应用段落样式
        self._apply_title_paragraph_format(para, h2_style)

        # 🔑 关键：为二级标题添加书签
        bookmark_name = self._make_heading_bookmark(number)
        self._add_bookmark_to_paragraph(para, bookmark_name)

    def _add_heading3(self, number, text):
        """
        添加三级标题
        :param number: 标题编号
        :param text: 标题文本
        """
        h3_style = self.style_manager.get_heading_style(3)
        para = self.doc.add_paragraph()

        # 编号（Times New Roman）
        num_run = para.add_run(f'{number} ')
        num_run.font.name = h3_style.get('number_font', 'Times New Roman')
        num_run.font.size = Pt(h3_style['size'])
        num_run.font.bold = h3_style.get('bold', False)

        # 标题文本（宋体）
        text_run = para.add_run(text)
        text_run.font.name = h3_style['font']
        text_run._element.rPr.rFonts.set(qn('w:eastAsia'), h3_style['font'])
        text_run.font.size = Pt(h3_style['size'])
        text_run.font.bold = h3_style.get('bold', False)

        # 应用段落样式
        self._apply_title_paragraph_format(para, h3_style)

        # 🔑 关键：为三级标题添加书签
        bookmark_name = self._make_heading_bookmark(number)
        self._add_bookmark_to_paragraph(para, bookmark_name)

    def _add_text_with_citations(self, paragraph, text, chinese_font, english_font, font_size, bold=False):
        """
        在段落中写入正文文本并处理参考文献引用
        """
        if not text:
            return

        last_index = 0
        for match in CITATION_PATTERN.finditer(text):
            plain_text = text[last_index:match.start()]
            if plain_text:
                self._append_text_run(paragraph, plain_text, chinese_font, english_font, font_size, bold)

            citation_number = int(match.group(1))
            self._append_citation_run(paragraph, citation_number, chinese_font, english_font, font_size, bold)
            last_index = match.end()

        if last_index < len(text):
            remaining = text[last_index:]
            if remaining:
                self._append_text_run(paragraph, remaining, chinese_font, english_font, font_size, bold)

    def _append_text_run(self, paragraph, text, chinese_font, english_font, font_size, bold=False):
        """向段落中添加普通文本"""
        if not text:
            return
        run = paragraph.add_run()
        self.style_manager.set_mixed_font(
            run,
            text,
            chinese_font=chinese_font,
            english_font=english_font,
            size=font_size,
            bold=bold
        )

    def _append_citation_run(self, paragraph, citation_number, chinese_font, english_font, font_size, bold=False):
        """插入形如[1]的引用并添加内部跳转"""
        citation_text = f'[{citation_number}]'
        target = self.reference_targets.get(citation_number)
        if not target:
            self._append_text_run(paragraph, citation_text, chinese_font, english_font, font_size, bold)
            return

        bookmark_name = None
        if citation_number not in self.reference_backlinks:
            bookmark_name = f'_Citation_{citation_number}'
            self.reference_backlinks[citation_number] = bookmark_name

        self._add_internal_reference_link(
            paragraph,
            citation_text,
            target['bookmark'],
            chinese_font,
            english_font,
            font_size,
            bold,
            bookmark_name_for_location=bookmark_name
        )

    def _add_internal_reference_link(self, paragraph, text, bookmark_name, chinese_font, english_font, font_size, bold=False, bookmark_name_for_location=None):
        """创建保持黑色字体的内部超链接（用于参考文献引用）"""
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        hyperlink.set(qn('w:history'), '1')

        run_element = OxmlElement('w:r')
        run_props = OxmlElement('w:rPr')

        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), english_font)
        fonts.set(qn('w:hAnsi'), english_font)
        fonts.set(qn('w:eastAsia'), chinese_font)
        run_props.append(fonts)

        size_element = OxmlElement('w:sz')
        size_element.set(qn('w:val'), str(int(font_size * 2)))
        run_props.append(size_element)

        if bold:
            bold_element = OxmlElement('w:b')
            bold_element.set(qn('w:val'), '1')
            run_props.append(bold_element)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        run_props.append(color)

        run_element.append(run_props)
        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)

        hyperlink.append(run_element)

        if bookmark_name_for_location:
            bookmark_id = self._get_next_bookmark_id()
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(bookmark_id))
            bookmark_start.set(qn('w:name'), bookmark_name_for_location)

            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(bookmark_id))

            p_element = paragraph._element
            p_element.append(bookmark_start)
            p_element.append(hyperlink)
            p_element.append(bookmark_end)
        else:
            paragraph._element.append(hyperlink)

        return hyperlink

    def _add_seq_field(self, paragraph, seq_type, prefix_text=''):
        """
        添加SEQ字段实现自动编号题注
        :param paragraph: 要添加字段的段落
        :param seq_type: 序列类型 'Figure' 或 'Table'
        :param prefix_text: 前缀文本（如 '图' 或 '表'）
        :return: paragraph 用于链式调用
        """
        if prefix_text:
            paragraph.add_run(prefix_text)

        run = paragraph.add_run()
        r = run._r

        # 字段开始标记
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar_begin)

        # 字段指令：SEQ Figure/Table \* ARABIC
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' SEQ {seq_type} \\* ARABIC '
        r.append(instrText)

        # 字段结束标记
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        r.append(fldChar_end)

        return paragraph

    def _add_chapter_based_seq_field(self, paragraph, seq_type, chapter_num, prefix_text=''):
        """
        添加基于章节的SEQ字段题注（如：图1-1, 表2-3）
        :param paragraph: 要添加字段的段落
        :param seq_type: 序列类型 'Figure' 或 'Table'
        :param chapter_num: 章节号
        :param prefix_text: 前缀文本（如 '图' 或 '表'）
        :return: paragraph 用于链式调用
        """
        if prefix_text:
            paragraph.add_run(prefix_text)

        # 添加章节号
        paragraph.add_run(str(chapter_num))
        paragraph.add_run('-')

        # 添加SEQ字段
        run = paragraph.add_run()
        r = run._r

        # 字段开始标记
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar_begin)

        # 字段指令：SEQ Figure/Table \* ARABIC \r 1（每章重置）
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' SEQ {seq_type}_{chapter_num} \\* ARABIC '
        r.append(instrText)

        # 字段结束标记
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        r.append(fldChar_end)

        return paragraph

    def _add_figure(self, figure_data):
        """
        添加图片（若图片缺失则跳过该图，避免占位符）
        :param figure_data: 图片数据
        """
        image_path = figure_data.get('path')
        if not image_path or not os.path.exists(image_path):
            return

        fig_style = self.style_manager.get_figure_style()

        # 添加空行（图与上文之间）
        self.doc.add_paragraph()

        # 插入真实图片
        img_para = self.doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            run = img_para.add_run()
            run.add_picture(image_path, width=Inches(5))
        except Exception as e:
            # 如果图片插入失败，显示错误信息但仍保留题注
            error_run = img_para.add_run(f'[图片加载失败: {str(e)}]')
            error_run.font.name = '宋体'
            error_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            error_run.font.size = Pt(10.5)

        # 图标题 - 使用SEQ字段实现自动编号
        caption_para = self.doc.add_paragraph()

        # 提取章节号（如 "1-1" 中的 "1"）
        number_str = figure_data.get('number', '1-1')
        chapter_num = number_str.split('-')[0] if '-' in number_str else '1'

        # 添加SEQ字段题注
        self._add_chapter_based_seq_field(caption_para, 'Figure', chapter_num, prefix_text='图')

        # 添加标题文本
        caption_text_run = caption_para.add_run(f" {figure_data.get('caption', '')}")
        caption_text_run.font.name = fig_style['caption']['font']
        caption_text_run._element.rPr.rFonts.set(qn('w:eastAsia'), fig_style['caption']['font'])
        caption_text_run.font.size = Pt(fig_style['caption']['size'])

        # 设置段落格式
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置题注段落中所有run的字体（包括SEQ字段部分）
        for run in caption_para.runs:
            run.font.name = fig_style['caption']['font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fig_style['caption']['font'])
            run.font.size = Pt(fig_style['caption']['size'])

        # 来源
        if figure_data.get('source'):
            source_para = self.doc.add_paragraph()
            source_run = source_para.add_run(f"来源：{figure_data['source']}")
            source_run.font.name = fig_style['source']['font']
            source_run._element.rPr.rFonts.set(qn('w:eastAsia'), fig_style['source']['font'])
            source_run.font.size = Pt(fig_style['source']['size'])

        # 添加空行（图与下文之间）
        self.doc.add_paragraph()

    def _add_table(self, table_data):
        """
        添加三线表
        :param table_data: 表格数据
        """
        tbl_style = self.style_manager.get_table_style()

        # 添加空行（表与上文之间）
        self.doc.add_paragraph()

        # 表标题 - 使用SEQ字段实现自动编号
        caption_para = self.doc.add_paragraph()

        # 提取章节号（如 "1-1" 中的 "1"）
        number_str = table_data.get('number', '1-1')
        chapter_num = number_str.split('-')[0] if '-' in number_str else '1'

        # 添加SEQ字段题注
        self._add_chapter_based_seq_field(caption_para, 'Table', chapter_num, prefix_text='表')

        # 添加标题文本
        caption_text_run = caption_para.add_run(f" {table_data.get('caption', '')}")
        caption_text_run.font.name = tbl_style['caption']['font']
        caption_text_run._element.rPr.rFonts.set(qn('w:eastAsia'), tbl_style['caption']['font'])
        caption_text_run.font.size = Pt(tbl_style['caption']['size'])

        # 设置段落格式
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置题注段落中所有run的字体（包括SEQ字段部分）
        for run in caption_para.runs:
            run.font.name = tbl_style['caption']['font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), tbl_style['caption']['font'])
            run.font.size = Pt(tbl_style['caption']['size'])

        # 防止标题与表格分页
        caption_para.paragraph_format.keep_with_next = True

        # 获取表格数据
        rows = table_data.get('rows', [])
        if not rows:
            return

        # 创建表格
        num_cols = len(rows[0]) if rows else 1
        num_rows = len(rows)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # 允许跨页断行
        self._allow_table_row_break(table)

        # 填充表格数据
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text

                # 设置单元格文字样式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = tbl_style['content_font']
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), tbl_style['content_font'])
                        run.font.size = Pt(tbl_style['content_size'])
                        # 表头加粗
                        if row_idx == 0:
                            run.font.bold = True

        # 表头在分页时重复
        self._repeat_table_header(table)

        # 设置三线表边框
        self._set_table_borders(table, tbl_style)

        # 来源
        if table_data.get('source'):
            source_para = self.doc.add_paragraph()
            source_run = source_para.add_run(f"来源：{table_data['source']}")
            source_run.font.name = tbl_style['source']['font']
            source_run._element.rPr.rFonts.set(qn('w:eastAsia'), tbl_style['source']['font'])
            source_run.font.size = Pt(tbl_style['source']['size'])

        # 添加空行（表与下文之间）
        self.doc.add_paragraph()

    def _set_table_borders(self, table, style_config):
        """
        设置三线表边框：上下 1.5pt，中线 0.5pt，去除竖线
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)

        # 清理并写入表级边框
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)

        top_border = style_config.get('top_border', 1.5)
        bottom_border = style_config.get('bottom_border', 1.5)
        middle_border = style_config.get('middle_border', 0.5)

        top_sz = self._border_size_value(top_border)
        bottom_sz = self._border_size_value(bottom_border)
        middle_sz = self._border_size_value(middle_border)

        borders_xml = f'''
            <w:tblBorders {nsdecls('w')}>
                <w:top w:val="single" w:sz="{top_sz}" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="{bottom_sz}" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="{middle_sz}" w:space="0" w:color="000000"/>
                <w:insideV w:val="nil"/>
            </w:tblBorders>
        '''
        tblPr.append(parse_xml(borders_xml))

        # 逐单元格写死边框，确保 Word 属性面板能读到正确值
        total_rows = len(table.rows)
        if total_rows == 0:
            return

        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            is_last = row_idx == total_rows - 1

            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                else:
                    for child in list(tcBorders):
                        tcBorders.remove(child)

                def _set_edge(edge_name, size_value):
                    edge = OxmlElement(f'w:{edge_name}')
                    if size_value is None:
                        edge.set(qn('w:val'), 'nil')
                    else:
                        edge.set(qn('w:val'), 'single')
                        edge.set(qn('w:sz'), size_value)
                        edge.set(qn('w:color'), '000000')
                        edge.set(qn('w:space'), '0')
                    tcBorders.append(edge)

                # 三线表无竖线
                _set_edge('left', None)
                _set_edge('right', None)

                if is_header:
                    _set_edge('top', top_sz)
                    # 单行表格直接使用底线，否则写中线
                    _set_edge('bottom', bottom_sz if total_rows == 1 else middle_sz)
                elif is_last:
                    _set_edge('top', None)
                    _set_edge('bottom', bottom_sz)
                else:
                    _set_edge('top', None)
                    _set_edge('bottom', None)

    def _border_size_value(self, points):
        """Word 边框宽度以 1/8pt 为单位"""
        return str(int(max(points, 0) * 8))

    def _allow_table_row_break(self, table):
        """允许表格行跨页断行"""
        for row in table.rows:
            tr = row._tr
            trPr = tr.trPr
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            cant_split = trPr.find(qn('w:cantSplit'))
            if cant_split is None:
                cant_split = OxmlElement('w:cantSplit')
                trPr.append(cant_split)
            cant_split.set(qn('w:val'), '0')

    def _repeat_table_header(self, table):
        """设置表头在分页时重复"""
        if not table.rows:
            return
        header_row = table.rows[0]
        tr = header_row._tr
        trPr = tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        tbl_header = trPr.find(qn('w:tblHeader'))
        if tbl_header is None:
            tbl_header = OxmlElement('w:tblHeader')
            trPr.append(tbl_header)
        tbl_header.set(qn('w:val'), 'on')

    def _text_to_omml(self, text):
        """
        将简单的数学文本转换为OMML格式
        支持下标（使用_）和基本运算符
        """
        # 设置Times New Roman字体和五号字体(10.5pt)
        font_name = 'Times New Roman'
        font_size = '21'  # 半磅单位，10.5pt * 2 = 21

        # 创建基本的文本运行
        def create_run(content, italic=False):
            sty = '<m:sty m:val="i"/>' if italic else ''
            return f'''<m:r>
                <m:rPr>{sty}</m:rPr>
                <w:rPr>
                    <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>
                    <w:sz w:val="{font_size}"/>
                </w:rPr>
                <m:t>{content}</m:t>
            </m:r>'''

        # 函数列表（不需要斜体）
        functions = {'sin', 'cos', 'tan', 'floor', 'log', 'ln', 'exp', 'max', 'min', 'ReLU', 'Concat'}

        result_parts = []
        i = 0

        while i < len(text):
            # 处理下标：x_i 或 word_num
            if i < len(text) - 2 and text[i].isalnum():
                # 查找下标
                j = i
                while j < len(text) and (text[j].isalnum() or text[j] == '_'):
                    if text[j] == '_':
                        break
                    j += 1

                if j < len(text) and text[j] == '_':
                    # 找到下标
                    base = text[i:j]
                    k = j + 1
                    while k < len(text) and text[k].isalnum():
                        k += 1
                    subscript = text[j+1:k]

                    # 检查base是否是函数名
                    is_func = base in functions
                    is_var = base.isalpha() and not is_func

                    result_parts.append(f'''<m:sSub>
                        <m:e>{create_run(base, italic=is_var)}</m:e>
                        <m:sub>{create_run(subscript, italic=subscript.isalpha())}</m:sub>
                    </m:sSub>''')
                    i = k
                    continue

            # 处理普通单词/变量
            if text[i].isalpha():
                j = i
                while j < len(text) and text[j].isalpha():
                    j += 1
                word = text[i:j]

                # 检查是否是函数名
                is_func = word in functions
                is_var = len(word) == 1 or not is_func

                result_parts.append(create_run(word, italic=is_var and not is_func))
                i = j
                continue

            # 处理数字和其他字符
            result_parts.append(create_run(text[i], italic=False))
            i += 1

        return ''.join(result_parts)

    def _add_formula(self, formula_data):
        """
        添加公式：使用OMML格式，公式居中，编号右对齐（同一行）
        :param formula_data: 公式数据
        """
        formula_style = self.style_manager.get_formula_style()
        formula_content = formula_data.get('content', '')
        formula_lines = [line.strip() for line in formula_content.split('\n') if line.strip()]
        if not formula_lines:
            return

        # 获取公式编号
        number_text = formula_data.get('number', '1-1')
        if not number_text.startswith('('):
            number_text = f'({number_text})'

        # 添加空行（公式与上文之间）
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.keep_with_next = True

        # 逐行处理公式
        for line_idx, line in enumerate(formula_lines):
            # 创建段落并强制 Word 识别为居中
            p = self.doc.add_paragraph()
            p.paragraph_format.keep_together = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 如果是最后一行，添加编号
            is_last_line = (line_idx == len(formula_lines) - 1)

            if is_last_line:
                # 最后一行：公式居中，编号右对齐（同一行）
                # 设置制表位：中央制表位和右对齐制表位
                self._add_tab_stop(p, position_cm=8.25, alignment='center')
                self._add_tab_stop(p, position_cm=16.0, alignment='right')

                # 添加制表符到中央位置
                p.add_run('\t')

            # 直接使用OxmlElement创建OMML结构
            try:
                # 创建oMathPara元素
                oMathPara = OxmlElement('m:oMathPara')

                # 创建oMath元素
                oMath = OxmlElement('m:oMath')

                # 解析公式并添加到oMath
                self._build_omml_runs(oMath, line, formula_style)

                # 将oMath添加到oMathPara
                oMathPara.append(oMath)

                # 根据配置写入公式段落对齐信息
                self._apply_math_justification(oMathPara, formula_style.get('alignment', 'center'))

                # 将oMathPara添加到段落
                p._element.append(oMathPara)

            except Exception as e:
                # 如果OMML构建失败，回退到文本模式
                print(f"OMML构建失败: {str(e)}, 使用文本模式")
                import traceback
                traceback.print_exc()
                print(f"公式内容: {line}")
                run = p.add_run(line)
                run.font.name = formula_style['font']
                run.font.size = Pt(formula_style['size'])

            if is_last_line:
                # 添加制表符到右侧位置
                p.add_run('\t')
                # 添加编号
                number_run = p.add_run(number_text)
                number_run.font.name = formula_style['font']
                number_run.font.size = Pt(formula_style['size'])

            # 保持段落在一起
            if line_idx < len(formula_lines) - 1:
                p.paragraph_format.keep_with_next = True

        # 添加空行（公式与下文之间）
        self.doc.add_paragraph()

    def _apply_math_justification(self, oMathPara, alignment):
        """在oMathPara上写入对齐信息，确保Word识别为真正居中"""
        oMathParaPr = oMathPara.find(qn('m:oMathParaPr'))
        if oMathParaPr is None:
            oMathParaPr = OxmlElement('m:oMathParaPr')
            oMathPara.insert(0, oMathParaPr)

        jc = oMathParaPr.find(qn('m:jc'))
        if jc is None:
            jc = OxmlElement('m:jc')
            oMathParaPr.append(jc)

        alignment_map = {
            'left': 'left',
            'center': 'center',
            'right': 'right',
            'centergroup': 'centerGroup'
        }
        jc.set(qn('m:val'), alignment_map.get(str(alignment).lower(), 'center'))

    def _build_omml_runs(self, oMath, text, formula_style):
        """
        构建OMML格式的run元素
        """
        font_name = formula_style.get('font', 'Times New Roman')
        font_size = str(int(formula_style.get('size', 10.5) * 2))  # 转换为半磅
        functions = {'sin', 'cos', 'tan', 'floor', 'log', 'ln', 'exp', 'max', 'min', 'ReLU', 'Concat'}

        i = 0
        while i < len(text):
            # 处理下标：x_i
            if i < len(text) - 2 and text[i].isalnum():
                j = i
                while j < len(text) and text[j].isalnum():
                    j += 1

                if j < len(text) and text[j] == '_':
                    # 找到下标
                    base = text[i:j]
                    k = j + 1
                    while k < len(text) and text[k].isalnum():
                        k += 1
                    subscript = text[j+1:k]

                    # 创建下标结构
                    sSub = OxmlElement('m:sSub')

                    # base元素
                    e = OxmlElement('m:e')
                    is_var = base.isalpha() and base not in functions
                    self._add_omml_text_run(e, base, font_name, font_size, italic=is_var)
                    sSub.append(e)

                    # subscript元素
                    sub = OxmlElement('m:sub')
                    self._add_omml_text_run(sub, subscript, font_name, font_size, italic=subscript.isalpha())
                    sSub.append(sub)

                    oMath.append(sSub)
                    i = k
                    continue

            # 处理普通单词
            if text[i].isalpha():
                j = i
                while j < len(text) and text[j].isalpha():
                    j += 1
                word = text[i:j]
                is_var = len(word) == 1 or word not in functions
                self._add_omml_text_run(oMath, word, font_name, font_size, italic=is_var and word not in functions)
                i = j
                continue

            # 其他字符
            self._add_omml_text_run(oMath, text[i], font_name, font_size, italic=False)
            i += 1

    def _add_omml_text_run(self, parent, text, font_name, font_size, italic=False):
        """添加一个OMML文本run"""
        r = OxmlElement('m:r')

        # 添加run属性
        rPr = OxmlElement('m:rPr')
        if italic:
            sty = OxmlElement('m:sty')
            sty.set(qn('m:val'), 'i')
            rPr.append(sty)
        r.append(rPr)

        # 添加Word格式属性
        w_rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        w_rPr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), font_size)
        w_rPr.append(sz)

        r.append(w_rPr)

        # 添加文本
        t = OxmlElement('m:t')
        t.text = text
        r.append(t)

        parent.append(r)

    def _set_header(self, title, section):
        """
        设置指定节的页眉
        :param title: 页眉文本（论文标题）
        :param section: 需要应用页眉的节
        """
        header_config = self.style_manager.config['body'].get('header')
        if not header_config:
            return

        header = section.header
        header.is_linked_to_previous = False
        self._clear_block_paragraphs(header)
        header_para = header.add_paragraph()

        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        header_para.alignment = alignment_map.get(header_config.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)

        header_run = header_para.add_run(title)
        font_name = header_config.get('font', '宋体')
        header_run.font.name = font_name
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        header_run.font.size = Pt(header_config.get('size', 9))
